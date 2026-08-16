"""
TeleRAG — Document Extractor

Extracts text from PDF/DOCX documents while preserving structure.
Focuses on preserving page numbers for citation accuracy.
"""
import logging
from pathlib import Path
from typing import Iterator

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)


class ExtractedPage:
    def __init__(self, page_number: int, text: str):
        self.page_number = page_number
        self.text = text


def extract_pdf(filepath: Path) -> Iterator[ExtractedPage]:
    """Extract text from a PDF, yielding ExtractedPage objects."""
    logger.info(f"Extracting PDF: {filepath.name}")
    try:
        with fitz.open(str(filepath)) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Use "text" dict or blocks if we need more layout control, 
                # but plain text is often fine for 3GPP specs if chunked cleanly.
                text = page.get_text("text")
                # 3GPP docs have header/footer, but for now we take full text.
                # Page numbers in 3GPP are usually physical page = printed page - offset.
                # We'll use the physical page number (1-indexed) as the citation page.
                yield ExtractedPage(page_number=page_num + 1, text=text)
    except Exception as e:
        logger.error(f"Error extracting PDF {filepath.name}: {e}")


def extract_docx(filepath: Path) -> Iterator[ExtractedPage]:
    """Extract text from a DOCX.
    
    DOCX has no strict physical pages. We'll simulate 'pages' by splitting
    text into chunks or treating the whole document as page 1.
    For this prototype, if it's DOCX, we just yield one big page, or break
    by headings. Since we prefer PDF for fidelity, this is a fallback.
    """
    logger.info(f"Extracting DOCX: {filepath.name}")
    try:
        from docx import Document
        doc = Document(str(filepath))
        text = "\n".join([p.text for p in doc.paragraphs])
        # Fallback: Treat as a single 'page' for extraction purposes
        yield ExtractedPage(page_number=1, text=text)
    except ImportError:
        logger.error("python-docx not installed. Cannot extract DOCX.")
    except Exception as e:
        logger.error(f"Error extracting DOCX {filepath.name}: {e}")


def extract_doc(filepath: Path) -> Iterator[ExtractedPage]:
    """Extract text from a legacy .doc (Word 97-2003) file.
    
    Uses aspose-words for pure Python cross-platform extraction of OLE2 docs.
    """
    logger.info(f"Extracting DOC: {filepath.name}")
    try:
        import aspose.words as aw
        import re
        
        doc = aw.Document(str(filepath))
        text = doc.get_text()
        
        # Remove Aspose evaluation watermark and any control characters
        text = re.sub(r'Evaluation Only.*?Aspose Pty Ltd\.', '', text, flags=re.DOTALL)
        text = re.sub(r'Created with an evaluation copy of Aspose\.Words.*?temporary-license/[^\r\n]*', '', text, flags=re.DOTALL)
        
        # Remove Word field codes like \x13 HYPERLINK ... \x14 ... \x15
        text = re.sub(r'\x13.*?\x14', '', text)
        text = re.sub(r'[\x13\x14\x15]', '', text)
        
        text = text.strip()
        
        if not text:
            logger.warning(f"Extracted text from {filepath.name} is empty.")
            
        yield ExtractedPage(page_number=1, text=text)
    except ImportError:
        logger.error("aspose-words not installed. Cannot extract .doc files.")


def extract_document(filepath: Path) -> list[ExtractedPage]:
    """Extract text from a supported document type."""
    ext = filepath.suffix.lower()
    pages = []
    
    with open(r'd:\mavenir\debug_extractor.txt', 'a') as f:
        f.write(f"extract_document called with {filepath}, ext: {ext}\n")
    
    logger.error(f"DEBUG: extract_document called with {filepath}, ext: {ext}")
    
    if ext == ".pdf":
        pages = list(extract_pdf(filepath))
    elif ext == ".docx":
        pages = list(extract_docx(filepath))
    elif ext == ".doc":
        try:
            pages = list(extract_doc(filepath))
        except Exception as e:
            with open(r'd:\mavenir\debug_extractor.txt', 'a') as f:
                f.write(f"Exception in extract_doc: {e}\n")
    else:
        logger.error(f"Unsupported extraction format: {ext}")
        
    with open(r'd:\mavenir\debug_extractor.txt', 'a') as f:
        f.write(f"extract_document returning {len(pages)} pages\n")
        
    logger.error(f"DEBUG: extract_document returning {len(pages)} pages")
    return pages
